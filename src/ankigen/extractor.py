"""Text extraction and vocabulary identification from PDFs, images, and Word documents."""

from __future__ import annotations

import base64
import logging
import os
import shutil
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Literal, NamedTuple, cast

from docx import Document
from openai import OpenAI
from pydantic import BaseModel, Field
from pypdf import PdfReader

from ankigen.anki_db import normalize_anki_term
from ankigen.chunking import estimate_tokens, split_text_for_extraction
from ankigen.cleaner import parse_hanja_token
from ankigen.extract_checkpoint import (
    ExtractRunCheckpoint,
    FileCheckpoint,
    clear_grammar_chunks,
    clear_vocab_chunks,
    source_changed,
)
from ankigen.llm import (
    LANGUAGE_CONFIG,
    PROVIDER_CONFIG,
    Language,
    format_llm_error,
    generate_structured_response,
    get_anthropic_client,
    get_extract_chunk_tokens,
    get_llm_max_output_tokens,
    get_model,
    get_provider,
    vocabulary_json_format_block,
)
from ankigen.models import GrammarItem
from ankigen.resume import durable_write

ExtractMode = Literal["vocab", "grammar", "all"]

logger = logging.getLogger("ankigen.extractor")

# Supported file extensions
PDF_EXTENSIONS = {".pdf"}
DOCX_EXTENSIONS = {".docx"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".webp"}
SUPPORTED_EXTENSIONS = PDF_EXTENSIONS | DOCX_EXTENSIONS | IMAGE_EXTENSIONS

# Default directories
DEFAULT_WATCH_DIR = "./watch"
DEFAULT_OUTPUT_DIR = "./inputs"
DEFAULT_PROCESSED_DIR = "./processed"


def get_watch_dir(lang: str | None = None) -> Path:
    """
    Get the watch directory from environment or default.

    Args:
        lang: If provided, returns language-specific watch dir

    Returns:
        Watch directory path
    """
    if lang:
        # Check for language-specific override first
        lang_env = f"ANKIGEN_WATCH_DIR_{lang.upper()}"
        lang_path = os.getenv(lang_env)
        if lang_path:
            return Path(lang_path)
        # Fall back to base/lang/
        return Path(os.getenv("ANKIGEN_WATCH_DIR", DEFAULT_WATCH_DIR)) / lang

    return Path(os.getenv("ANKIGEN_WATCH_DIR", DEFAULT_WATCH_DIR))


def get_output_dir() -> Path:
    """Get the output directory from environment or default."""
    return Path(os.getenv("ANKIGEN_OUTPUT_DIR", DEFAULT_OUTPUT_DIR))


def get_processed_dir(lang: str | None = None) -> Path:
    """
    Get the processed files directory from environment or default.

    Args:
        lang: If provided, returns language-specific processed dir

    Returns:
        Processed directory path
    """
    if lang:
        # Check for language-specific override first
        lang_env = f"ANKIGEN_PROCESSED_DIR_{lang.upper()}"
        lang_path = os.getenv(lang_env)
        if lang_path:
            return Path(lang_path)
        # Fall back to base/lang/
        return Path(os.getenv("ANKIGEN_PROCESSED_DIR", DEFAULT_PROCESSED_DIR)) / lang

    return Path(os.getenv("ANKIGEN_PROCESSED_DIR", DEFAULT_PROCESSED_DIR))


class VocabularyResponse(BaseModel):
    """Response model for vocabulary extraction."""

    words: list[str] = Field(
        ...,
        description="List of vocabulary words extracted from the text",
    )


def extract_text_from_pdf(path: Path) -> str:
    """
    Extract text content from a PDF file.

    Args:
        path: Path to the PDF file

    Returns:
        Extracted text content
    """
    try:
        file_size = path.stat().st_size / 1024  # KB
        logger.debug("Processing PDF: %s (%.1f KB)", path.name, file_size)
    except OSError:
        logger.debug("Processing PDF: %s", path.name)
    logger.info("Extracting text from PDF: %s", path.name)

    start_time = time.time()
    reader = PdfReader(path)

    text_parts: list[str] = []
    for page_num, page in enumerate(reader.pages, 1):
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
            logger.debug("Page %d: extracted %d characters", page_num, len(page_text))

    full_text = "\n\n".join(text_parts)
    elapsed = time.time() - start_time
    logger.debug("PDF extraction completed in %.2fs", elapsed)
    logger.info("Extracted %d characters from %d pages", len(full_text), len(reader.pages))
    return full_text


def extract_text_from_docx(path: Path, *, with_headings: bool = False) -> str:
    """
    Extract text content from a Word document (.docx).

    Args:
        path: Path to the .docx file
        with_headings: If True, prefix each Heading 1/2/3 paragraph with
            ``[H1] ``/``[H2] ``/``[H3] `` markers so downstream LLM prompts
            can use document structure as a signal. Body paragraphs and
            tables are unchanged.

    Returns:
        Extracted text content
    """
    try:
        file_size = path.stat().st_size / 1024  # KB
        logger.debug("Processing DOCX: %s (%.1f KB)", path.name, file_size)
    except OSError:
        logger.debug("Processing DOCX: %s", path.name)
    logger.info("Extracting text from DOCX: %s", path.name)

    start_time = time.time()
    doc = Document(str(path))

    text_parts: list[str] = []
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        if with_headings:
            style_name = para.style.name if para.style is not None else ""
            if style_name.startswith("Heading 1"):
                text_parts.append(f"[H1] {para.text}")
            elif style_name.startswith("Heading 2"):
                text_parts.append(f"[H2] {para.text}")
            elif style_name.startswith("Heading 3"):
                text_parts.append(f"[H3] {para.text}")
            else:
                text_parts.append(para.text)
        else:
            text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    full_text = "\n".join(text_parts)
    elapsed = time.time() - start_time
    logger.debug("DOCX extraction completed in %.2fs", elapsed)
    logger.info("Extracted %d characters from DOCX", len(full_text))
    return full_text


def extract_text_from_image(path: Path, lang: Language = "zh") -> str:
    """
    Extract text from an image using GPT-4 Vision OCR.

    Args:
        path: Path to the image file
        lang: Language of the text in the image

    Returns:
        Extracted text content
    """
    try:
        file_size = path.stat().st_size / 1024  # KB
        logger.debug("Processing image: %s (%.1f KB)", path.name, file_size)
    except OSError:
        logger.debug("Processing image: %s", path.name)
    logger.info("Performing OCR on image: %s", path.name)

    # Read and encode the image
    with open(path, "rb") as f:
        image_data = base64.b64encode(f.read()).decode("utf-8")

    # Determine media type
    suffix = path.suffix.lower()
    media_types = {
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".webp": "image/webp",
    }
    media_type = media_types.get(suffix, "image/png")
    logger.debug("Image type: %s, base64 size: %d bytes", media_type, len(image_data))

    # Use raw OpenAI-compatible client for vision where supported.
    # Anthropic uses its native SDK path below.
    provider = get_provider()
    config = PROVIDER_CONFIG[provider]
    base_url = os.getenv("LLM_BASE_URL") or config["base_url"]
    api_key = os.getenv("LLM_API_KEY", "")

    # Build headers for OpenRouter if needed
    default_headers = {}
    if provider == "openrouter" or "openrouter.ai" in base_url:
        site_url = os.getenv("OPENROUTER_SITE_URL")
        app_name = os.getenv("OPENROUTER_APP_NAME")
        if site_url:
            default_headers["HTTP-Referer"] = site_url
        if app_name:
            default_headers["X-Title"] = app_name

    lang_name = LANGUAGE_CONFIG[lang]["name"]

    # Use configured model, with vision-capable fallback per provider.
    if provider == "openrouter":
        model = os.getenv("LLM_VISION_MODEL") or "openai/gpt-4o"
    elif provider == "anthropic":
        model = os.getenv("LLM_VISION_MODEL") or "claude-sonnet-4-6"
    else:
        model = os.getenv("LLM_VISION_MODEL") or "gpt-4o"

    logger.debug("Calling %s for %s OCR via %s", model, lang_name, provider)

    start_time = time.time()
    if provider == "anthropic":
        anthropic_client = get_anthropic_client()
        anthropic_content: list[dict[str, Any]] = [
            {
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": media_type,
                    "data": image_data,
                },
            },
            {
                "type": "text",
                "text": f"Extract all {lang_name} text from this image.",
            },
        ]
        anthropic_response = anthropic_client.messages.create(
            model=model,
            max_tokens=get_llm_max_output_tokens(),
            system=(
                f"You are an OCR assistant. Extract all {lang_name} text from the image. "
                "Preserve the original text exactly as it appears. "
                "Do not translate or interpret the text, just transcribe it."
            ),
            messages=[
                {
                    "role": "user",
                    "content": cast(Any, anthropic_content),
                }
            ],
        )
        extracted_text = "\n".join(
            block.text for block in anthropic_response.content if block.type == "text"
        ).strip()
    else:
        openai_client = OpenAI(
            base_url=base_url,
            api_key=api_key,
            default_headers=default_headers if default_headers else None,
        )
        openai_response = openai_client.chat.completions.create(
            model=model,
            messages=[
                {
                    "role": "system",
                    "content": f"You are an OCR assistant. Extract all {lang_name} text from the image. "
                    "Preserve the original text exactly as it appears. "
                    "Do not translate or interpret the text, just transcribe it.",
                },
                {
                    "role": "user",
                    "content": [
                        {
                            "type": "image_url",
                            "image_url": {"url": f"data:{media_type};base64,{image_data}"},
                        },
                        {
                            "type": "text",
                            "text": f"Extract all {lang_name} text from this image.",
                        },
                    ],
                },
            ],
            max_tokens=get_llm_max_output_tokens(),
        )
        extracted_text = openai_response.choices[0].message.content or ""

    elapsed = time.time() - start_time
    logger.debug("OCR completed in %.2fs", elapsed)
    logger.info("Extracted %d characters from image", len(extracted_text))
    return extracted_text


def _merge_vocab_words(chunks_words: list[list[str]], lang: Language) -> list[str]:
    """Dedupe vocab across chunks, preserving order and Hanja annotations.

    For Korean, the dedupe key is the bare word (without any ``(漢字)``
    annotation); if any chunk produced a ``한글(漢字)`` form for the same bare
    word, that annotated form replaces a previously-seen bare entry.
    """
    if lang != "ko":
        seen: dict[str, str] = {}
        order: list[str] = []
        for chunk in chunks_words:
            for w in chunk:
                if w not in seen:
                    seen[w] = w
                    order.append(w)
        return order

    merged: dict[str, str] = {}
    ko_order: list[str] = []
    for chunk in chunks_words:
        for w in chunk:
            bare, hanja = parse_hanja_token(w)
            if bare not in merged:
                merged[bare] = w
                ko_order.append(bare)
            elif hanja and not parse_hanja_token(merged[bare])[1]:
                # Upgrade a previously bare entry with an annotated one.
                merged[bare] = w
    return [merged[k] for k in ko_order]


def _vocab_from_checkpoint(
    run_checkpoint: ExtractRunCheckpoint,
    file_entry: FileCheckpoint,
    lang: Language,
) -> list[str] | None:
    """Rebuild merged vocab from saved chunk JSONL when vocab pass already finished."""
    if file_entry.status not in ("vocab_done", "grammar_done"):
        return None
    by_index = run_checkpoint.load_all_vocab_chunks(file_entry)
    if not by_index:
        return None
    ordered = [by_index[i] for i in sorted(by_index)]
    return _merge_vocab_words(ordered, lang)


def identify_vocabulary(
    text: str,
    lang: Language = "zh",
    *,
    run_checkpoint: ExtractRunCheckpoint | None = None,
    file_entry: FileCheckpoint | None = None,
) -> list[str]:
    """
    Identify vocabulary words from text using LLM.

    Long inputs are automatically split into chunks (each under
    ``get_extract_chunk_tokens()``) so a single call never exceeds the
    provider's per-minute input-token budget. Results are merged with order
    preserved and duplicates removed; Korean Hanja annotations win over bare
    forms when both appear across chunks.

    Args:
        text: The text to extract vocabulary from
        lang: Language of the text

    Returns:
        List of vocabulary words
    """
    logger.debug("Identifying vocabulary from %d characters", len(text))

    model = get_model()
    lang_name = LANGUAGE_CONFIG[lang]["name"]

    hanja_rule = (
        " EXCEPTION for Korean: if the source text presents a Sino-Korean word "
        "together with its Hanja form (e.g. '음식(飮食)', '한자(漢字)'), keep that "
        "exact '한글(漢字)' annotation in the returned word so downstream tools "
        "can reuse the Hanja without a separate lookup. Do not invent Hanja "
        "annotations that are not in the source text."
        if lang == "ko"
        else ""
    )

    system_prompt = (
        f"You are a {lang_name} language expert. "
        "Extract important vocabulary words from the given text. "
        "Focus on words that would be useful for language learners: "
        "nouns, verbs, adjectives, adverbs, and common expressions. "
        "Exclude very common words (like 'the', 'is', 'a' equivalents). "
        "Return each word in its dictionary/base form. "
        f"IMPORTANT: Return ONLY the {lang_name} characters/script. "
        "Do NOT include any romanization (pinyin, romaja, etc.), "
        "pronunciation guides, or parenthetical annotations."
        + hanja_rule
        + "\n\n"
        + vocabulary_json_format_block(lang)
    )

    chunk_limit = get_extract_chunk_tokens()
    chunks = split_text_for_extraction(text, chunk_limit)
    total_est = estimate_tokens(text)
    if len(chunks) > 1:
        logger.info(
            "Splitting vocab extract into %d chunks (~%d est. tokens total, max %d tokens/chunk)",
            len(chunks),
            total_est,
            chunk_limit,
        )

    start_time = time.time()
    chunks_words: list[list[str]] = []
    for idx, chunk in enumerate(chunks):
        chunk_num = idx + 1
        chunk_est = estimate_tokens(chunk)
        cached: list[str] | None = None
        if run_checkpoint is not None and file_entry is not None:
            cached = run_checkpoint.load_vocab_chunk(file_entry, idx)

        if cached is not None:
            logger.info(
                "Resuming vocab chunk %d/%d (%d words cached, model=%s)",
                chunk_num,
                len(chunks),
                len(cached),
                model,
            )
            chunks_words.append(cached)
            continue

        logger.info(
            "LLM vocab chunk %d/%d (~%d est. tokens, model=%s)",
            chunk_num,
            len(chunks),
            chunk_est,
            model,
        )
        chunk_start = time.time()
        response = generate_structured_response(
            response_model=VocabularyResponse,
            system_prompt=system_prompt,
            user_prompt=(
                f"Extract vocabulary words from this {lang_name} text and respond in JSON:\n\n{chunk}"
            ),
        )
        chunk_words = list(response.words)  # type: ignore[attr-defined]
        chunks_words.append(chunk_words)
        logger.info(
            "LLM vocab chunk %d/%d finished in %.2fs → %d words",
            chunk_num,
            len(chunks),
            time.time() - chunk_start,
            len(chunk_words),
        )
        if run_checkpoint is not None and file_entry is not None:
            run_checkpoint.save_vocab_chunk(file_entry, idx, chunk_words)

    merged = _merge_vocab_words(chunks_words, lang)
    elapsed = time.time() - start_time
    logger.info(
        "Identified %d vocabulary words in %.2fs (%d chunk(s))",
        len(merged),
        elapsed,
        len(chunks),
    )
    if run_checkpoint is not None and file_entry is not None:
        run_checkpoint.mark_vocab_done(file_entry)
    return merged


def get_file_type(path: Path) -> str:
    """
    Determine the file type based on extension.

    Returns:
        'pdf', 'docx', 'image', or raises ValueError for unsupported types
    """
    suffix = path.suffix.lower()
    if suffix in PDF_EXTENSIONS:
        return "pdf"
    elif suffix in DOCX_EXTENSIONS:
        return "docx"
    elif suffix in IMAGE_EXTENSIONS:
        return "image"
    else:
        raise ValueError(
            f"Unsupported file type: {suffix}. "
            f"Supported: PDF ({', '.join(PDF_EXTENSIONS)}), "
            f"Word ({', '.join(DOCX_EXTENSIONS)}), "
            f"Images ({', '.join(IMAGE_EXTENSIONS)})"
        )


def extract_source_text(
    path: Path,
    lang: Language = "zh",
    *,
    with_headings: bool = False,
) -> str:
    """
    Extract raw text from a PDF, DOCX, or image file.

    Centralised so vocab and grammar pipelines can share a single text-extraction
    pass over the same file (cheaper for OCR especially).

    Args:
        path: Source file.
        lang: Content language (only used for OCR prompts).
        with_headings: When the source is a DOCX, include [H1]/[H2]/[H3] markers
            so LLM prompts can detect document structure. No effect on PDF/images.

    Returns:
        Extracted text. May be empty if the file has no extractable text.
    """
    file_type = get_file_type(path)
    if file_type == "pdf":
        return extract_text_from_pdf(path)
    elif file_type == "docx":
        return extract_text_from_docx(path, with_headings=with_headings)
    else:
        return extract_text_from_image(path, lang)


def extract_vocabulary_from_file(path: Path, lang: Language = "zh") -> list[str]:
    """
    Extract vocabulary words from a PDF, Word document, or image file.

    This is the main entry point that:
    1. Detects file type
    2. Extracts text (PDF/DOCX text extraction or OCR for images)
    3. Identifies vocabulary using LLM

    Args:
        path: Path to the PDF, DOCX, or image file
        lang: Language of the content

    Returns:
        List of vocabulary words
    """
    text = extract_source_text(path, lang)

    if not text.strip():
        logger.warning("No text extracted from %s", path)
        return []

    return identify_vocabulary(text, lang)


def get_supported_files(directory: Path, *, recursive: bool = False) -> list[Path]:
    """
    Get all supported files (PDFs, Word docs, and images) from a directory.

    Excludes hidden files (starting with .) and temporary files (starting with ~).

    Args:
        directory: Directory to scan
        recursive: If True, walk into subdirectories as well

    Returns:
        List of paths to supported files
    """
    if not directory.exists():
        return []

    iterator = directory.rglob("*") if recursive else directory.iterdir()

    files: list[Path] = []
    for path in iterator:
        if path.name.startswith(".") or path.name.startswith("~"):
            continue
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS:
            files.append(path)

    return sorted(files)


class FolderResult(NamedTuple):
    """Outcome of a folder/watch run, per pipeline."""

    vocab_path: Path | None
    grammar_path: Path | None
    num_files: int


def _move_files(files: list[Path], processed_dir: Path) -> None:
    """Move ``files`` into ``processed_dir`` with conflict-safe renaming."""
    if not files:
        return
    processed_dir.mkdir(parents=True, exist_ok=True)
    for file_path in files:
        dest = processed_dir / file_path.name
        if dest.exists():
            stem = file_path.stem
            suffix = file_path.suffix
            counter = 1
            while dest.exists():
                dest = processed_dir / f"{stem}_{counter}{suffix}"
                counter += 1
        shutil.move(str(file_path), str(dest))
        logger.debug("Moved %s to %s", file_path.name, dest)
    logger.info("Moved %d files to %s", len(files), processed_dir)


def _write_vocab_output(unique_words: list[str], output_path: Path) -> None:
    """Write vocab words, append+dedupe if file already exists."""
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if output_path.exists():
        with open(output_path, encoding="utf-8") as f:
            existing_words = {line.strip() for line in f if line.strip()}
        new_words = [w for w in unique_words if w not in existing_words]
        if new_words:
            logger.info(
                "Appending %d new words to %s (skipping %d duplicates)",
                len(new_words),
                output_path,
                len(unique_words) - len(new_words),
            )
            with open(output_path, "a", encoding="utf-8") as f:
                for word in new_words:
                    f.write(word + "\n")
                durable_write(f)
        else:
            logger.info("All words already exist in %s", output_path)
    else:
        with open(output_path, "w", encoding="utf-8") as f:
            for word in unique_words:
                f.write(word + "\n")
            durable_write(f)
    logger.info("Vocab output written to %s", output_path)


def process_folder(
    lang: Language = "zh",
    *,
    source_dir: Path | None = None,
    mode: ExtractMode = "vocab",
    move_processed: bool | None = None,
    recursive: bool = False,
    exclude_words: set[str] | None = None,
    exclude_patterns: set[str] | None = None,
    use_checkpoint: bool = True,
    fresh: bool = False,
) -> FolderResult:
    """
    Process all supported files from a directory (watch folder or ad-hoc).

    Routing by ``mode``:

    - ``vocab`` — produces ``{output_dir}/{lang}/{YYYYMMDD}.txt`` (append+dedupe).
    - ``grammar`` — produces ``{output_dir}/{lang}/{YYYYMMDD}_grammar.jsonl``
      (append+dedupe by NFC-normalised pattern).
    - ``all`` — runs both pipelines, sharing one text-extraction pass per file
      (cheaper for OCR especially). Files are only counted as "processed" if
      both passes succeed.

    Move semantics (default):

    - ``move_processed=None`` → move iff ``mode == "all"``. So plain ``vocab``
      and ``grammar`` runs leave files in place (you may want to run the other
      pass on them next).
    - Pass ``move_processed=True/False`` to override.

    Args:
        lang: Content language (also picks watch / processed sub-folder).
        source_dir: Directory to scan. ``None`` → use the configured watch dir
            for this language.
        mode: Which pipeline(s) to run.
        move_processed: Override move behaviour. Default depends on ``mode``.
        recursive: When ``True``, walk subdirectories of ``source_dir``.
        exclude_words: NFC-normalised words to skip in the vocab pass.
        exclude_patterns: NFC-normalised patterns to skip in the grammar pass.
        use_checkpoint: When True, write staging checkpoints and incremental outputs.
        fresh: When True, ignore existing staging for this run key and start clean.
    """
    # Lazy import to avoid an extractor↔grammar circular dependency at import time.
    from ankigen.extract_checkpoint import ExtractRunCheckpoint, init_manifest
    from ankigen.grammar import (
        extract_grammar_items,
        write_grammar_jsonl,
    )

    if move_processed is None:
        move_processed = mode == "all"

    if source_dir is None:
        source_dir = get_watch_dir(lang)
    output_dir = get_output_dir()
    processed_dir = get_processed_dir(lang)

    files = get_supported_files(source_dir, recursive=recursive)
    if not files:
        logger.info("No files found in %s", source_dir)
        return FolderResult(None, None, 0)

    logger.info("Found %d files to process in %s (mode=%s)", len(files), source_dir, mode)

    today = datetime.now().strftime("%Y%m%d")
    vocab_output = output_dir / lang / f"{today}.txt" if mode in ("vocab", "all") else None
    grammar_output = (
        output_dir / lang / f"{today}_grammar.jsonl" if mode in ("grammar", "all") else None
    )

    run_checkpoint: ExtractRunCheckpoint | None = None
    if use_checkpoint:
        manifest = init_manifest(
            lang=lang,
            mode=mode,
            source_dir=source_dir,
            date=today,
            file_paths=files,
            fresh=fresh,
        )
        run_checkpoint = ExtractRunCheckpoint(manifest)
        done, total = run_checkpoint.count_resumable(mode)
        if done:
            logger.info("Resuming extract run (%d/%d files already complete)", done, total)

    all_words: list[str] = []
    all_grammar_items: list[GrammarItem] = []
    processed_files: list[Path] = []
    num_files = len(files)

    for file_index, file_path in enumerate(files, start=1):
        try:
            file_entry: FileCheckpoint | None = None
            if run_checkpoint is not None:
                from ankigen.extract_checkpoint import find_file_entry

                file_entry = find_file_entry(run_checkpoint.manifest, file_path)
                if file_entry and run_checkpoint.should_skip_file(file_entry, mode):
                    logger.info(
                        "Skipping file %d/%d (already complete): %s",
                        file_index,
                        num_files,
                        file_path.name,
                    )
                    processed_files.append(file_path)
                    continue

            logger.info("Processing file %d/%d: %s", file_index, num_files, file_path.name)

            text: str | None = None
            if run_checkpoint is not None and file_entry is not None:
                if source_changed(file_entry, file_path):
                    clear_vocab_chunks(file_entry, run_checkpoint)
                    clear_grammar_chunks(file_entry, run_checkpoint)
                    file_entry.status = "pending"
                    file_entry.vocab_chunks = 0
                    file_entry.grammar_chunks = 0
                else:
                    text = run_checkpoint.load_cached_text(file_entry, file_path)

            if text is None:
                text = extract_source_text(
                    file_path,
                    lang,
                    with_headings=(mode in ("grammar", "all")),
                )
                if run_checkpoint is not None and file_entry is not None and text.strip():
                    run_checkpoint.save_text(file_entry, file_path, text)

            if not text.strip():
                logger.warning("No text extracted from %s", file_path.name)
                continue

            logger.info(
                "Extracted %d characters (~%d est. tokens) from %s",
                len(text),
                estimate_tokens(text),
                file_path.name,
            )

            file_succeeded = True

            if mode in ("vocab", "all"):
                try:
                    words: list[str] | None = None
                    if (
                        run_checkpoint is not None
                        and file_entry is not None
                        and file_entry.status in ("vocab_done", "grammar_done")
                    ):
                        words = _vocab_from_checkpoint(run_checkpoint, file_entry, lang)
                        if words is not None:
                            logger.info(
                                "Reusing %d cached vocab words for %s",
                                len(words),
                                file_path.name,
                            )
                    if words is None:
                        words = identify_vocabulary(
                            text,
                            lang,
                            run_checkpoint=run_checkpoint,
                            file_entry=file_entry,
                        )
                    all_words.extend(words)
                    logger.info("Extracted %d words from %s", len(words), file_path.name)
                    if vocab_output is not None:
                        _write_vocab_output(words, vocab_output)
                except Exception as exc:
                    err = format_llm_error(exc)
                    logger.error("Vocab extraction failed for %s: %s", file_path.name, err)
                    if run_checkpoint is not None and file_entry is not None:
                        run_checkpoint.mark_failed(file_entry, err)
                    file_succeeded = False

            if mode in ("grammar", "all"):
                try:
                    from ankigen.grammar import _grammar_from_checkpoint

                    items: list[GrammarItem] | None = None
                    if (
                        run_checkpoint is not None
                        and file_entry is not None
                        and file_entry.status == "grammar_done"
                    ):
                        items = _grammar_from_checkpoint(run_checkpoint, file_entry)
                        if items is not None:
                            logger.info(
                                "Reusing %d cached grammar item(s) for %s",
                                len(items),
                                file_path.name,
                            )
                    if items is None:
                        items = extract_grammar_items(
                            text,
                            lang,
                            run_checkpoint=run_checkpoint,
                            file_entry=file_entry,
                        )
                    all_grammar_items.extend(items)
                    logger.info("Extracted %d grammar item(s) from %s", len(items), file_path.name)
                    if grammar_output is not None:
                        write_grammar_jsonl(items, grammar_output, append=True)
                except Exception as exc:
                    err = format_llm_error(exc)
                    logger.error("Grammar extraction failed for %s: %s", file_path.name, err)
                    if run_checkpoint is not None and file_entry is not None:
                        run_checkpoint.mark_failed(file_entry, err)
                    file_succeeded = False

            if file_succeeded:
                processed_files.append(file_path)
        except Exception as exc:  # belt-and-braces around the whole per-file block
            logger.error("Failed to process %s: %s", file_path.name, format_llm_error(exc))

    if run_checkpoint is not None and run_checkpoint.all_files_complete(mode):
        run_checkpoint.mark_run_complete()

    if mode in ("vocab", "all"):
        if all_words:
            seen: set[str] = set()
            unique_words: list[str] = []
            for word in all_words:
                if word not in seen:
                    unique_words.append(word)
                    seen.add(word)

            logger.info(
                "Total: %d unique words from %d files",
                len(unique_words),
                len(processed_files),
            )

            if exclude_words:
                before = len(unique_words)
                unique_words = [
                    w for w in unique_words if normalize_anki_term(w) not in exclude_words
                ]
                skipped = before - len(unique_words)
                if skipped:
                    logger.info(
                        "Skipped %d words already present in Anki (%d remaining)",
                        skipped,
                        len(unique_words),
                    )

            if vocab_output is not None and not vocab_output.exists():
                _write_vocab_output(unique_words, vocab_output)
            elif vocab_output is not None:
                logger.info("Vocab output already written incrementally to %s", vocab_output)
        else:
            logger.warning("No vocabulary extracted from any files")

    if mode in ("grammar", "all"):
        if all_grammar_items:
            if exclude_patterns:
                before = len(all_grammar_items)
                all_grammar_items = [
                    it
                    for it in all_grammar_items
                    if normalize_anki_term(it.pattern) not in exclude_patterns
                ]
                skipped = before - len(all_grammar_items)
                if skipped:
                    logger.info("Skipped %d grammar pattern(s) already present in Anki", skipped)

            if grammar_output is not None and not grammar_output.exists():
                write_grammar_jsonl(all_grammar_items, grammar_output, append=False)
            elif grammar_output is not None:
                logger.info(
                    "Grammar output already written incrementally to %s",
                    grammar_output,
                )
        else:
            logger.warning("No grammar items extracted from any files")

    if move_processed and processed_files:
        _move_files(processed_files, processed_dir)
    elif processed_files and not move_processed:
        logger.info(
            "Leaving %d processed file(s) in place (move disabled for mode=%s)",
            len(processed_files),
            mode,
        )

    return FolderResult(vocab_output, grammar_output, len(processed_files))


def process_watch_folder(
    lang: Language = "zh",
    *,
    move_processed: bool = True,
    exclude_words: set[str] | None = None,
) -> tuple[Path | None, int]:
    """
    Backwards-compatible wrapper around :func:`process_folder`.

    Preserves the original signature (vocab-only, watch dir, returns
    ``(output_path, num_files)``) so external callers and tests built against
    the previous API keep working. Defaults to vocab mode and the configured
    watch folder.
    """
    result = process_folder(
        lang=lang,
        source_dir=None,
        mode="vocab",
        move_processed=move_processed,
        exclude_words=exclude_words,
    )
    return result.vocab_path, result.num_files
